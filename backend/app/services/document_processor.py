import PyPDF2
import pdfplumber
from docx import Document
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging
import re

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Document processing service"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt
        }
    
    async def process_document(
        self,
        document_path: Path,
        file_id: str
    ) -> Dict[str, Any]:
        try:
            suffix = document_path.suffix.lower()
            
            if suffix not in self.supported_formats:
                raise ValueError(f"Unsupported document format: {suffix}")
            
            logger.info(f"Processing document: {document_path}")
            
            # Process based on file type
            processor = self.supported_formats[suffix]
            content = await processor(document_path)
            
            result = {
                'file_id': file_id,
                'filename': document_path.name,
                'format': suffix,
                'text': content['text'],
                'metadata': content.get('metadata', {}),
                'structure': content.get('structure', {})
            }
            
            logger.info(f"Document processed: {len(content['text'])} characters")
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            raise
    
    async def _process_pdf(self, pdf_path: Path) -> Dict[str, Any]:
        """Extract text from PDF"""
        text_parts = []
        metadata = {}
        
        try:
            # Use pdfplumber for better text extraction
            with pdfplumber.open(pdf_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                
                for page in pdf.pages:
                    if page_text:
                        text_parts.append(page_text)
            
            full_text = '\n\n'.join(text_parts)
            
            return {
                'text': full_text,
                'metadata': metadata,
                'structure': {'type': 'pdf', 'pages': metadata['pages']}
            }
            
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise
    
    async def _process_docx(self, docx_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX"""
        try:
            doc = Document(docx_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    tables_text.append(row_text)
            
            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            if tables_text:
                full_text += '\n\n' + '\n'.join(tables_text)
            
            metadata = {
                'paragraphs': len(paragraphs),
                'tables': len(doc.tables)
            }
            
            return {
                'text': full_text,
                'metadata': metadata,
                'structure': {'type': 'docx', 'has_tables': len(doc.tables) > 0}
            }
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    async def _process_txt(self, txt_path: Path) -> Dict[str, Any]:
        """Extract text from TXT file"""
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            # Count lines
            lines = text.split('\n')
            
            metadata = {
                'lines': len(lines),
                'encoding': 'utf-8'
            }
            
            return {
                'text': text,
                'metadata': metadata,
                'structure': {'type': 'txt'}
            }
            
        except Exception as e:
            logger.error(f"TXT processing failed: {e}")
            raise
    
    async def extract_template_fields(
        self,
        template_path: Path
    ) -> List[str]:
        """
        Extract template field placeholders from document
        Looks for patterns like {{field_name}} or {field_name}
        
        Args:
            template_path: Path to template document
        
        Returns:
            List of field names found in template
        """
        try:
            # Process the template document
            content = await self.process_document(template_path, "template")
            text = content['text']
            
            # Find all placeholders (supports {{field}} and {field} patterns)
            pattern = r'\{\{([^}]+)\}\}|\{([^}]+)\}'
            matches = re.findall(pattern, text)
            
            # Flatten and clean field names
            fields = []
            for match in matches:
                field = match[0] if match[0] else match[1]
                field = field.strip()
                if field and field not in fields:
                    fields.append(field)
            
            logger.info(f"Found {len(fields)} template fields: {fields}")
            return fields
            
        except Exception as e:
            logger.error(f"Template field extraction failed: {e}")
            raise
    
    async def batch_process(
        self,
        documents: List[tuple[Path, str]]
    ) -> List[Dict[str, Any]]:
        """
        Process multiple documents
        
        Args:
            documents: List of (document_path, file_id) tuples
        
        Returns:
            List of processed document results
        """
        results = []
        for doc_path, file_id in documents:
            try:
                result = await self.process_document(doc_path, file_id)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {doc_path}: {e}")
                # Continue with other documents
        
        return results


# Global instance
document_processor = DocumentProcessor()

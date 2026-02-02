from pathlib import Path
from typing import Dict, Any, List, Optional
import logging
from datetime import datetime
import uuid
from jinja2 import Template, Environment, FileSystemLoader
from docx import Document
import re

from app.config import settings
from app.models.schemas import CompletedDocument, TemplateField

logger = logging.getLogger(__name__)


class TemplateFiller:
    """Document template completion service"""
    
    def __init__(self):
        self.template_dir = settings.template_dir
        self.output_dir = settings.output_dir
        self.jinja_env = Environment(loader=FileSystemLoader(str(self.template_dir)))
    
    async def fill_template(
        self,
        template_path: Path,
        data: Dict[str, Any],
        output_name: Optional[str] = None
    ) -> CompletedDocument:
        try:
            document_id = str(uuid.uuid4())
            logger.info(f"Filling template: {template_path.name}")
            
            # Determine template type and process accordingly
            suffix = template_path.suffix.lower()
            
            if suffix == '.docx':
                output_path = await self._fill_docx_template(
                    template_path, data, document_id, output_name
                )
            elif suffix == '.txt':
                output_path = await self._fill_text_template(
                    template_path, data, document_id, output_name
                )
            else:
                raise ValueError(f"Unsupported template format: {suffix}")
            
            # Create field list
            fields = [
                TemplateField(
                    field_name=key,
                    field_value=value,
                    field_type=type(value).__name__
                )
                for key, value in data.items()
            ]
            
            completed_doc = CompletedDocument(
                document_id=document_id,
                template_name=template_path.name,
                fields=fields,
                output_path=str(output_path),
                created_at=datetime.now()
            )
            
            logger.info(f"Template completed: {output_path}")
            return completed_doc
            
        except Exception as e:
            logger.error(f"Template filling failed: {e}")
            raise
    
    async def _fill_docx_template(
        self,
        template_path: Path,
        data: Dict[str, Any],
        document_id: str,
        output_name: Optional[str] = None
    ) -> Path:
        """Fill DOCX template"""
        
        # Load template document
        doc = Document(template_path)
        
        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_placeholders_in_text(paragraph, data)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders_in_text(paragraph, data)
        
        # Save completed document
        if output_name:
            output_filename = output_name
        else:
            output_filename = f"completed_{document_id}.docx"
        
        output_path = self.output_dir / output_filename
        doc.save(str(output_path))
        
        return output_path
    
    def _replace_placeholders_in_text(
        self,
        paragraph,
        data: Dict[str, Any]
    ):
        """Replace placeholders in paragraph text"""
        
        # Get full paragraph text
        full_text = paragraph.text
        
        # Find and replace placeholders (supports {{field}} and {field})
        for key, value in data.items():
            # Replace {{key}} pattern
            full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
            # Replace {key} pattern
            full_text = full_text.replace(f"{{{key}}}", str(value))
        
        # Update paragraph text while preserving formatting
        if full_text != paragraph.text:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ''
            
            # Add new text to first run
            if paragraph.runs:
                paragraph.runs[0].text = full_text
            else:
                paragraph.add_run(full_text)
    
    async def _fill_text_template(
        self,
        template_path: Path,
        data: Dict[str, Any],
        document_id: str,
        output_name: Optional[str] = None
    ) -> Path:
        """Fill text template using Jinja2"""
        
        # Read template
        with open(template_path, 'r', encoding='utf-8') as f:
            template_content = f.read()
        
        # Create Jinja2 template
        template = Template(template_content)
        
        # Render with data
        rendered_content = template.render(**data)
        
        # Save completed document
        if output_name:
            output_filename = output_name
        else:
            output_filename = f"completed_{document_id}.txt"
        
        output_path = self.output_dir / output_filename
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(rendered_content)
        
        return output_path
    
    async def batch_fill_templates(
        self,
        templates: List[Path],
        data: Dict[str, Any]
    ) -> List[CompletedDocument]:
        """
        Fill multiple templates with the same data
        
        Args:
            templates: List of template paths
            data: Data to fill templates with
        
        Returns:
            List of completed documents
        """
        results = []
        for template_path in templates:
            try:
                result = await self.fill_template(template_path, data)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to fill template {template_path}: {e}")
                # Continue with other templates
        
        return results
    
    async def preview_template(
        self,
        template_path: Path,
        sample_data: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        Generate preview of template with sample data
        
        Args:
            template_path: Path to template
            sample_data: Optional sample data (uses placeholders if None)
        
        Returns:
            Preview text
        """
        try:
            suffix = template_path.suffix.lower()
            
            if suffix == '.txt':
                with open(template_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                if sample_data:
                    template = Template(content)
                    return template.render(**sample_data)
                return content
            
            elif suffix == '.docx':
                doc = Document(template_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                return '\n\n'.join(paragraphs)
            
            else:
                return f"Preview not available for {suffix} files"
                
        except Exception as e:
            logger.error(f"Template preview failed: {e}")
            return f"Error generating preview: {e}"


# Global instance
template_filler = TemplateFiller()
